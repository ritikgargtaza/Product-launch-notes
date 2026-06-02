#!/usr/bin/env python3
"""
Convert a launch-notes markdown file to .docx.

Usage: python3 md_to_docx.py <path/to/launch-notes.md>
Writes <path/to/launch-notes.docx> alongside the source.

Tailored to the launch-notes structure produced by install/launch-notes.md:
- # / ## / ### headings
- **bold** and `code` inline runs
- > blockquotes
- | pipe tables |
- - bullet lists
- --- horizontal rules
- Plain paragraphs
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


INLINE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")


def shade_cell(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def add_inline_runs(paragraph, text):
    """Split text on inline markup and add runs with appropriate formatting."""
    if not text:
        return
    parts = INLINE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Menlo"
            run.font.size = Pt(10)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def is_table_row(line):
    return line.strip().startswith("|") and line.strip().endswith("|")


def is_table_separator(line):
    s = line.strip()
    return s.startswith("|") and set(s.replace("|", "").replace(" ", "")) <= set("-:")


def parse_table_row(line):
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return cells


def add_table(doc, rows, has_header):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Light Grid Accent 1"
    table.autofit = True
    for i, row in enumerate(rows):
        for j in range(cols):
            cell = table.rows[i].cells[j]
            cell.text = ""
            content = row[j] if j < len(row) else ""
            para = cell.paragraphs[0]
            add_inline_runs(para, content)
            if has_header and i == 0:
                for run in para.runs:
                    run.bold = True
                shade_cell(cell, "F2F2F2")


def convert(md_path: Path):
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    # Base style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Slightly tighter margins for a one-pager feel
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Horizontal rule — skip (visual breaks come from headings/spacing)
        if stripped == "---":
            i += 1
            continue

        # Blank line
        if not stripped:
            i += 1
            continue

        # Headings
        if stripped.startswith("### "):
            p = doc.add_heading(level=3)
            add_inline_runs(p, stripped[4:])
            i += 1
            continue
        if stripped.startswith("## "):
            p = doc.add_heading(level=2)
            add_inline_runs(p, stripped[3:])
            i += 1
            continue
        if stripped.startswith("# "):
            p = doc.add_heading(level=1)
            add_inline_runs(p, stripped[2:])
            i += 1
            continue

        # Blockquote (single-line block — call-out style)
        if stripped.startswith(">"):
            content = stripped.lstrip("> ").strip()
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # Table
        if is_table_row(line):
            rows = []
            has_header = False
            # Look ahead: header row + separator + body rows
            rows.append(parse_table_row(line))
            j = i + 1
            if j < len(lines) and is_table_separator(lines[j]):
                has_header = True
                j += 1
            while j < len(lines) and is_table_row(lines[j]):
                rows.append(parse_table_row(lines[j]))
                j += 1
            add_table(doc, rows, has_header)
            doc.add_paragraph()  # spacing after table
            i = j
            continue

        # Bullet
        if stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, stripped[2:])
            i += 1
            continue

        # Plain paragraph — may span multiple non-blank lines
        para_lines = [stripped]
        j = i + 1
        while j < len(lines):
            nxt = lines[j].strip()
            if not nxt:
                break
            if (nxt.startswith("#") or nxt.startswith(">") or nxt.startswith("- ")
                    or nxt.startswith("* ") or nxt == "---" or is_table_row(lines[j])):
                break
            para_lines.append(nxt)
            j += 1
        p = doc.add_paragraph()
        add_inline_runs(p, " ".join(para_lines))
        i = j

    out_path = md_path.with_suffix(".docx")
    doc.save(out_path)
    print(f"wrote {out_path}")


def main():
    if len(sys.argv) != 2:
        print("usage: md_to_docx.py <launch-notes.md>", file=sys.stderr)
        sys.exit(2)
    md_path = Path(sys.argv[1]).expanduser().resolve()
    if not md_path.exists():
        print(f"file not found: {md_path}", file=sys.stderr)
        sys.exit(1)
    convert(md_path)


if __name__ == "__main__":
    main()
